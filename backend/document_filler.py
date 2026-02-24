"""
Document Filler — Fill uploaded documents with answers from the Knowledge Base.

Reads questions from an uploaded file (xlsx, docx, txt, csv), searches the KB
for each question, gets LLM answers with citations, and writes answers back
into a new file in the same format.
"""

import os
import csv
import json
import time
from io import StringIO

import vector_store
from config import Config


# ---------------------------------------------------------------------------
# Per-question answering (reusable across formats)
# ---------------------------------------------------------------------------

def answer_question(question: str, llm_client, kb_source_names: set = None) -> dict:
    """
    Search KB for a question, get LLM answer with citation.
    Returns { answer, source, kb_chunks_used }.
    """
    question = question.strip()
    if not question or len(question) < 5:
        return {'answer': '', 'source': '', 'kb_chunks_used': 0}

    # Search KB
    hits = vector_store.search(question, top_k=6)
    hits = [h for h in hits if h.get('filename') and h['filename'] != 'unknown' and h.get('doc_id', -1) != -1]

    if hits:
        # Build context from KB chunks
        context_parts = []
        sources = set()
        for i, h in enumerate(hits[:4]):
            src = h['filename']
            sources.add(src)
            context_parts.append(f"[Source: {src}]\n{h['text']}")
        context_str = "\n\n".join(context_parts)
        sources_list = ", ".join(sorted(sources))

        prompt = f"""Answer the following question using ONLY the provided context.
Cite the source document name at the end. Use ONLY these exact source names: {sources_list}

Context:
{context_str}

Question: {question}

Provide a concise answer (2-4 sentences max). End with the source in format [Source: exact_filename]
If the context doesn't contain the answer, say so and answer from your general knowledge, citing as [Source: External Knowledge].

Answer:"""
    else:
        # No KB hits — use LLM general knowledge
        prompt = f"""Answer the following question concisely (2-4 sentences max) using your general knowledge.
End your answer with [Source: External Knowledge]

Question: {question}

Answer:"""

    answer = llm_client.invoke(prompt)
    
    # Determine source from answer text
    source = 'External Knowledge'
    if hits:
        for h in hits:
            if h['filename'] in answer:
                source = h['filename']
                break

    return {
        'answer': answer.strip(),
        'source': source,
        'kb_chunks_used': len(hits),
    }


# ---------------------------------------------------------------------------
# Format-specific fillers
# ---------------------------------------------------------------------------

def _fill_xlsx(input_path: str, output_path: str, llm_client) -> dict:
    """Fill an xlsx file: read questions from column A, write answers in column B, sources in column C."""
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill

    wb = openpyxl.load_workbook(input_path)
    total_questions = 0
    total_answered = 0
    qa_pairs = []

    for ws in wb.worksheets:
        # Find the last row with data
        max_row = ws.max_row or 0
        if max_row == 0:
            continue

        # Add header for Answer and Source columns if they don't exist
        if ws.cell(1, 1).value and not ws.cell(1, 2).value:
            ws.cell(1, 2, 'Answer')
            ws.cell(1, 3, 'Source')
            ws.cell(1, 2).font = Font(bold=True, color='FFFFFF')
            ws.cell(1, 3).font = Font(bold=True, color='FFFFFF')
            ws.cell(1, 2).fill = PatternFill(start_color='8B5CF6', end_color='8B5CF6', fill_type='solid')
            ws.cell(1, 3).fill = PatternFill(start_color='8B5CF6', end_color='8B5CF6', fill_type='solid')
            start_row = 2
        else:
            start_row = 1

        for row_idx in range(start_row, max_row + 1):
            question = str(ws.cell(row_idx, 1).value or '').strip()
            if not question or len(question) < 5:
                continue

            total_questions += 1
            result = answer_question(question, llm_client)
            if result['answer']:
                answer_text = result['answer']
                ws.cell(row_idx, 2, answer_text)
                ws.cell(row_idx, 2).alignment = Alignment(wrap_text=True)
                ws.cell(row_idx, 3, result['source'])
                total_answered += 1
                qa_pairs.append({'question': question, 'answer': answer_text, 'source': result['source']})

        ws.column_dimensions['B'].width = 80
        ws.column_dimensions['C'].width = 40

    wb.save(output_path)
    return {'total_questions': total_questions, 'total_answered': total_answered, 'qa_pairs': qa_pairs}


def _fill_docx(input_path: str, output_path: str, llm_client) -> dict:
    """Fill a docx file: find question paragraphs, insert answer paragraphs below."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document(input_path)
    total_questions = 0
    total_answered = 0
    qa_pairs = []

    paragraphs_to_process = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text and len(text) > 10 and text.endswith('?'):
            paragraphs_to_process.append((i, text))

    for idx, question in reversed(paragraphs_to_process):
        total_questions += 1
        result = answer_question(question, llm_client)
        if result['answer']:
            total_answered += 1
            answer_para = doc.paragraphs[idx]._element

            new_p = doc.add_paragraph()
            new_p.text = f"Answer: {result['answer']}"
            for run in new_p.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

            src_p = doc.add_paragraph()
            src_p.text = f"[Source: {result['source']}]"
            for run in src_p.runs:
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                run.font.italic = True

            answer_para.addnext(src_p._element)
            answer_para.addnext(new_p._element)
            qa_pairs.append({'question': question, 'answer': result['answer'], 'source': result['source']})

    # Reverse qa_pairs since we processed in reverse
    qa_pairs.reverse()
    doc.save(output_path)
    return {'total_questions': total_questions, 'total_answered': total_answered, 'qa_pairs': qa_pairs}


def _fill_txt(input_path: str, output_path: str, llm_client) -> dict:
    """Fill a txt file: find question lines, append answers below."""
    with open(input_path, 'r', encoding='utf-8', errors='replace') as f:
        lines = f.readlines()

    total_questions = 0
    total_answered = 0
    output_lines = []
    qa_pairs = []

    for line in lines:
        output_lines.append(line.rstrip('\n'))
        text = line.strip()
        if text and len(text) > 10 and text.endswith('?'):
            total_questions += 1
            result = answer_question(text, llm_client)
            if result['answer']:
                total_answered += 1
                output_lines.append(f"  → {result['answer']}")
                output_lines.append(f"  [Source: {result['source']}]")
                output_lines.append('')
                qa_pairs.append({'question': text, 'answer': result['answer'], 'source': result['source']})

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(output_lines))

    return {'total_questions': total_questions, 'total_answered': total_answered, 'qa_pairs': qa_pairs}


def _fill_csv(input_path: str, output_path: str, llm_client) -> dict:
    """Fill a csv file: add Answer and Source columns."""
    with open(input_path, 'r', encoding='utf-8', errors='replace') as f:
        reader = csv.reader(f)
        rows = list(reader)

    total_questions = 0
    total_answered = 0
    output_rows = []
    qa_pairs = []

    for i, row in enumerate(rows):
        if not row:
            output_rows.append(row)
            continue

        question = row[0].strip()
        if i == 0 and question.lower() in ('question', 'questions', 'q', ''):
            output_rows.append(row + ['Answer', 'Source'])
            continue

        if question and len(question) > 10:
            total_questions += 1
            result = answer_question(question, llm_client)
            if result['answer']:
                total_answered += 1
                output_rows.append(row + [result['answer'], result['source']])
                qa_pairs.append({'question': question, 'answer': result['answer'], 'source': result['source']})
            else:
                output_rows.append(row + ['', ''])
        else:
            output_rows.append(row + ['', ''])

    with open(output_path, 'w', encoding='utf-8', newline='') as f:
        writer = csv.writer(f)
        writer.writerows(output_rows)

    return {'total_questions': total_questions, 'total_answered': total_answered, 'qa_pairs': qa_pairs}


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

SUPPORTED_FORMATS = {'.xlsx', '.xls', '.docx', '.txt', '.csv'}

def fill_document(input_path: str, filename: str, llm_client) -> dict:
    """
    Fill a document with answers. Returns { output_path, output_filename, stats }.
    """
    ext = os.path.splitext(filename)[1].lower()

    if ext not in SUPPORTED_FORMATS:
        raise ValueError(f"Unsupported format: {ext}. Supported: {', '.join(SUPPORTED_FORMATS)}")

    # Create output path
    output_dir = os.path.join(Config.UPLOAD_FOLDER, '_filled')
    os.makedirs(output_dir, exist_ok=True)
    ts = time.strftime('%Y%m%d_%H%M%S')
    base_name = os.path.splitext(filename)[0]
    output_filename = f"{base_name}_filled_{ts}{ext}"
    output_path = os.path.join(output_dir, output_filename)

    # Dispatch to format handler
    if ext in ('.xlsx', '.xls'):
        stats = _fill_xlsx(input_path, output_path, llm_client)
    elif ext == '.docx':
        stats = _fill_docx(input_path, output_path, llm_client)
    elif ext == '.txt':
        stats = _fill_txt(input_path, output_path, llm_client)
    elif ext == '.csv':
        stats = _fill_csv(input_path, output_path, llm_client)
    else:
        raise ValueError(f"Unsupported format: {ext}")

    print(f"📝 Filled document: {output_filename} — {stats['total_answered']}/{stats['total_questions']} questions answered")

    return {
        'output_path': output_path,
        'output_filename': output_filename,
        'stats': stats,
        'qa_pairs': stats.get('qa_pairs', []),
    }
